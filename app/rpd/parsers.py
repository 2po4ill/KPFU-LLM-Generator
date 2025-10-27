"""
Multi-format RPD document parsers for KPFU LLM Generator
"""

import os
import logging
from typing import Dict, Any, Optional, List, Union
from pathlib import Path
import chardet

# Document processing imports
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    import magic
except ImportError:
    magic = None

logger = logging.getLogger(__name__)


class RPDParsingError(Exception):
    """Custom exception for RPD parsing errors"""
    pass


class FileTypeDetector:
    """Detects file types for RPD documents"""
    
    @staticmethod
    def detect_file_type(file_path: Union[str, Path]) -> str:
        """
        Detect file type based on extension and content
        
        Args:
            file_path: Path to the file
            
        Returns:
            File type string ('pdf', 'docx', 'xlsx', 'unknown')
        """
        file_path = Path(file_path)
        
        # First check by extension
        extension = file_path.suffix.lower()
        extension_map = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'doc',
            '.xlsx': 'xlsx',
            '.xls': 'xls'
        }
        
        if extension in extension_map:
            return extension_map[extension]
        
        # Fallback to magic number detection if available
        if magic and file_path.exists():
            try:
                mime_type = magic.from_file(str(file_path), mime=True)
                mime_map = {
                    'application/pdf': 'pdf',
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
                    'application/msword': 'doc',
                    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
                    'application/vnd.ms-excel': 'xls'
                }
                return mime_map.get(mime_type, 'unknown')
            except Exception as e:
                logger.warning(f"Magic detection failed: {e}")
        
        return 'unknown'


class BaseRPDParser:
    """Base class for RPD document parsers"""
    
    def __init__(self):
        self.supported_extensions = []
    
    def can_parse(self, file_path: Union[str, Path]) -> bool:
        """Check if this parser can handle the file"""
        file_type = FileTypeDetector.detect_file_type(file_path)
        return file_type in self.supported_extensions
    
    def parse(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Parse RPD document and extract structured data
        
        Args:
            file_path: Path to the RPD document
            
        Returns:
            Dictionary with extracted RPD data
        """
        raise NotImplementedError("Subclasses must implement parse method")
    
    def _detect_encoding(self, file_path: Union[str, Path]) -> str:
        """Detect file encoding for text files"""
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read(10000)  # Read first 10KB
                result = chardet.detect(raw_data)
                return result.get('encoding', 'utf-8')
        except Exception:
            return 'utf-8'


class PDFRPDParser(BaseRPDParser):
    """Parser for PDF RPD documents using PyPDF2"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['pdf']
        
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF parsing")
    
    def parse(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Parse PDF RPD document
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dictionary with extracted text and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise RPDParsingError(f"File not found: {file_path}")
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract metadata
                metadata = {}
                if pdf_reader.metadata:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', ''),
                        'creation_date': pdf_reader.metadata.get('/CreationDate', ''),
                    }
                
                # Extract text from all pages
                text_content = []
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content.append({
                                'page': page_num + 1,
                                'text': page_text.strip()
                            })
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                
                return {
                    'file_type': 'pdf',
                    'file_path': str(file_path),
                    'metadata': metadata,
                    'pages': len(pdf_reader.pages),
                    'content': text_content,
                    'raw_text': '\n'.join([page['text'] for page in text_content])
                }
                
        except Exception as e:
            raise RPDParsingError(f"Failed to parse PDF file {file_path}: {e}")


class WordRPDParser(BaseRPDParser):
    """Parser for Word RPD documents using python-docx"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['docx', 'doc']
        
        if Document is None:
            raise ImportError("python-docx is required for Word document parsing")
    
    def parse(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Parse Word RPD document
        
        Args:
            file_path: Path to the Word file
            
        Returns:
            Dictionary with extracted text and structure
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise RPDParsingError(f"File not found: {file_path}")
        
        try:
            doc = Document(file_path)
            
            # Extract document properties
            core_props = doc.core_properties
            metadata = {
                'title': core_props.title or '',
                'author': core_props.author or '',
                'subject': core_props.subject or '',
                'created': str(core_props.created) if core_props.created else '',
                'modified': str(core_props.modified) if core_props.modified else '',
            }
            
            # Extract paragraphs
            paragraphs = []
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    paragraphs.append({
                        'index': i,
                        'text': paragraph.text.strip(),
                        'style': paragraph.style.name if paragraph.style else 'Normal'
                    })
            
            # Extract tables
            tables = []
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                tables.append({
                    'index': table_idx,
                    'data': table_data,
                    'rows': len(table_data),
                    'columns': len(table_data[0]) if table_data else 0
                })
            
            return {
                'file_type': 'docx',
                'file_path': str(file_path),
                'metadata': metadata,
                'paragraphs': paragraphs,
                'tables': tables,
                'raw_text': '\n'.join([p['text'] for p in paragraphs])
            }
            
        except Exception as e:
            raise RPDParsingError(f"Failed to parse Word file {file_path}: {e}")


class ExcelRPDParser(BaseRPDParser):
    """Parser for Excel RPD documents using pandas"""
    
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['xlsx', 'xls']
        
        if pd is None:
            raise ImportError("pandas is required for Excel parsing")
    
    def parse(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Parse Excel RPD document
        
        Args:
            file_path: Path to the Excel file
            
        Returns:
            Dictionary with extracted data from all sheets
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise RPDParsingError(f"File not found: {file_path}")
        
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            sheets_data = {}
            
            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    # Convert to structured format
                    sheet_data = {
                        'name': sheet_name,
                        'shape': df.shape,
                        'columns': df.columns.tolist(),
                        'data': df.fillna('').to_dict('records'),
                        'raw_text': df.fillna('').to_string()
                    }
                    
                    sheets_data[sheet_name] = sheet_data
                    
                except Exception as e:
                    logger.warning(f"Failed to parse sheet '{sheet_name}': {e}")
            
            # Combine all text content
            all_text = []
            for sheet_data in sheets_data.values():
                all_text.append(sheet_data['raw_text'])
            
            return {
                'file_type': 'xlsx',
                'file_path': str(file_path),
                'sheets': sheets_data,
                'sheet_names': list(sheets_data.keys()),
                'raw_text': '\n\n'.join(all_text)
            }
            
        except Exception as e:
            raise RPDParsingError(f"Failed to parse Excel file {file_path}: {e}")


class RPDParserFactory:
    """Factory class for creating appropriate RPD parsers"""
    
    def __init__(self):
        self.parsers = {
            'pdf': PDFRPDParser,
            'docx': WordRPDParser,
            'doc': WordRPDParser,
            'xlsx': ExcelRPDParser,
            'xls': ExcelRPDParser,
        }
    
    def get_parser(self, file_path: Union[str, Path]) -> BaseRPDParser:
        """
        Get appropriate parser for the file
        
        Args:
            file_path: Path to the file
            
        Returns:
            Parser instance
            
        Raises:
            RPDParsingError: If no suitable parser found
        """
        file_type = FileTypeDetector.detect_file_type(file_path)
        
        if file_type not in self.parsers:
            raise RPDParsingError(f"Unsupported file type: {file_type}")
        
        try:
            return self.parsers[file_type]()
        except ImportError as e:
            raise RPDParsingError(f"Parser dependencies not available: {e}")
    
    def parse_rpd(self, file_path: Union[str, Path]) -> Dict[str, Any]:
        """
        Parse RPD document using appropriate parser
        
        Args:
            file_path: Path to the RPD document
            
        Returns:
            Dictionary with extracted RPD data
        """
        parser = self.get_parser(file_path)
        return parser.parse(file_path)
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats"""
        return list(self.parsers.keys())


# Global parser factory instance
rpd_parser_factory = RPDParserFactory()


def parse_rpd_document(file_path: Union[str, Path]) -> Dict[str, Any]:
    """
    Convenience function to parse RPD document
    
    Args:
        file_path: Path to the RPD document
        
    Returns:
        Dictionary with extracted RPD data
    """
    return rpd_parser_factory.parse_rpd(file_path)