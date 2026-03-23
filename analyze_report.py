import docx

doc = docx.Document('Девличаров Е.Т. Отчет (1).docx')

print(f"📄 Document Statistics:")
print(f"   Total paragraphs: {len(doc.paragraphs)}")
print(f"   Total tables: {len(doc.tables)}")
print(f"   Total sections: {len(doc.sections)}")

print("\n" + "="*80)
print("📋 DOCUMENT STRUCTURE ANALYSIS")
print("="*80)

# Collect headings and structure
structure = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    
    style = para.style.name
    
    # Track headings and important sections
    if 'Heading' in style or any(keyword in text for keyword in ['Введение', 'Заключение', 'Содержание', 'Список использованных']):
        structure.append({
            'index': i,
            'style': style,
            'text': text[:200]
        })

print("\n🔍 Main Sections and Headings:")
for item in structure[:30]:
    indent = "  " if 'Heading 2' in item['style'] else ""
    print(f"{indent}[{item['style']}] {item['text']}")

# Analyze content sections
print("\n" + "="*80)
print("📊 CONTENT BREAKDOWN")
print("="*80)

sections = {
    'Введение': [],
    'Разработка расширенной системы': [],
    'Доработка ИИ противника': [],
    'Обучение игрока': [],
    'Меню настроек': [],
    'Заключение': [],
    'Список источников': []
}

current_section = None
for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    
    if 'Введение' in text and len(text) < 50:
        current_section = 'Введение'
    elif 'Разработка расширенной системы' in text:
        current_section = 'Разработка расширенной системы'
    elif 'Доработка ИИ противника' in text:
        current_section = 'Доработка ИИ противника'
    elif 'Обучение игрока' in text:
        current_section = 'Обучение игрока'
    elif 'меню настроек' in text.lower():
        current_section = 'Меню настроек'
    elif 'Заключение' in text and len(text) < 50:
        current_section = 'Заключение'
    elif 'Список использованных источников' in text:
        current_section = 'Список источников'
    
    if current_section and len(sections[current_section]) < 5:
        sections[current_section].append(text[:150])

for section, content in sections.items():
    if content:
        print(f"\n📌 {section}:")
        for line in content[:3]:
            print(f"   • {line}")

# Table analysis
if doc.tables:
    print("\n" + "="*80)
    print("📊 TABLES FOUND")
    print("="*80)
    for i, table in enumerate(doc.tables):
        print(f"\nTable {i+1}: {len(table.rows)} rows × {len(table.columns)} columns")
        if len(table.rows) > 0:
            first_row = [cell.text.strip()[:30] for cell in table.rows[0].cells]
            print(f"   First row: {first_row}")

print("\n" + "="*80)
print("✅ Analysis Complete")
print("="*80)
