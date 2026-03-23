import docx

doc = docx.Document('Девличаров Е.Т. Отчет (1).docx')

print("="*80)
print("🎯 CORE CONTENT EXTRACTION")
print("="*80)

# Extract Introduction
print("\n" + "="*80)
print("📖 ВВЕДЕНИЕ (Introduction)")
print("="*80)
intro_started = False
intro_content = []
for para in doc.paragraphs:
    text = para.text.strip()
    if 'Введение' == text and para.style.name == 'Heading 1':
        intro_started = True
        continue
    if intro_started:
        if para.style.name == 'Heading 1':
            break
        if text:
            intro_content.append(text)

for line in intro_content[:15]:
    print(line)

# Extract main technical sections
sections_to_extract = [
    ('Разработка расширенной системы персонажей и оружия', 'CHARACTER_SYSTEM'),
    ('Доработка ИИ противника', 'AI_SYSTEM'),
    ('Обучение игрока и прогрессия', 'PROGRESSION_SYSTEM'),
    ('Разработка меню настроек и аудио менеджера', 'UI_AUDIO_SYSTEM')
]

for section_title, section_key in sections_to_extract:
    print("\n" + "="*80)
    print(f"🔧 {section_title.upper()}")
    print("="*80)
    
    section_started = False
    section_content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if section_title in text and para.style.name == 'Heading 1':
            section_started = True
            continue
        
        if section_started:
            if para.style.name == 'Heading 1' and text and len(text) < 100:
                break
            if text and len(text) > 20:
                section_content.append(text)
                if len(section_content) >= 10:
                    break
    
    for i, line in enumerate(section_content[:10], 1):
        print(f"{i}. {line[:200]}{'...' if len(line) > 200 else ''}")

# Extract Conclusion
print("\n" + "="*80)
print("🎓 ЗАКЛЮЧЕНИЕ (Conclusion)")
print("="*80)
conclusion_started = False
conclusion_content = []
for para in doc.paragraphs:
    text = para.text.strip()
    if 'Заключение' == text and para.style.name == 'Heading 1':
        conclusion_started = True
        continue
    if conclusion_started:
        if para.style.name == 'Heading 1':
            break
        if text:
            conclusion_content.append(text)

for line in conclusion_content[:10]:
    print(line)

# Extract competencies table
print("\n" + "="*80)
print("📊 КОМПЕТЕНЦИИ (Competencies)")
print("="*80)
if doc.tables:
    table = doc.tables[0]
    print(f"\nТаблица компетенций ({len(table.rows)} строк):\n")
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        if i == 0:
            print(f"{'Шифр':<15} | Описание")
            print("-" * 80)
        else:
            print(f"{cells[0]:<15} | {cells[1][:60]}...")

print("\n" + "="*80)
print("✅ Core Content Extraction Complete")
print("="*80)
